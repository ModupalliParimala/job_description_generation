"""Save the JD"""

import pypandoc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

DOCS_FOLDER = "docs"


def save_jd_and_retrieve(llm_response, job_title):
    """Persist the JD into folder"""

    save_jd_txt(llm_response, file_name=job_title)
    save_jd_doc(llm_response, file_name=job_title)
    save_jd_pdf(llm_response, file_name=job_title)

    return job_title


def save_jd_doc(llm_response, file_name):
    """Persist the JD as Document"""
    bullet_style = "List Bullet"

    doc = Document()
    page_title = doc.add_heading(
        file_name,
        level=0,
    )
    page_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Job Description:", level=1)
    doc.add_paragraph(llm_response.get("Description"))

    doc.add_heading("Responsibilities:", level=1)
    for responsibility in llm_response.get("Responsibilities"):
        doc.add_paragraph(responsibility, bullet_style)

    doc.add_heading("Skills:", level=1)
    for responsibility in llm_response.get("Skills"):
        doc.add_paragraph(responsibility, bullet_style)

    doc.add_heading("Experience:", level=1)
    for responsibility in llm_response.get("Experience"):
        doc.add_paragraph(responsibility, bullet_style)

    for _ in range(2):
        doc.add_paragraph()

    paragraph = doc.add_paragraph(llm_response.get("Closing Statement"))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.runs[0]  # Access the first run in the paragraph
    # run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.italic = True

    doc.save(f"{DOCS_FOLDER}/{file_name}.docx")
    return f"{file_name}.docx"


def save_jd_pdf(llm_response, file_name):
    """Persist the JD as PDF"""
    c = canvas.Canvas(f"{DOCS_FOLDER}/{file_name}.pdf", pagesize=A4)
    width, height = A4

    # Set title of the document
    c.setFont("Helvetica-Bold", 18)
    c.drawCentredString(width / 2.0, height - 100, f"{file_name}")

    # Description
    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, height - 150, f"Description:")

    c.setFont("Helvetica-Oblique", 12) 
    y_position = height - 170
    wrapped_lines = wrap_text(c,llm_response.get("Description"), width - 150)
    c.drawString(100, y_position,"")
    for line in wrapped_lines:
        c.drawString(115, y_position, line)  # Indent text after bullet
        y_position -= 15  # Move down for the next line

    # Responsibilities Header
    y_position -= 20
    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, y_position, "Responsibilities:")

    # Responsibilities List
    c.setFont("Helvetica-Oblique", 12) 
    bullet = "•"
    responsibilities = llm_response.get("Responsibilities")
    for idx, responsibility in enumerate(responsibilities, 1):
        c.drawString(120, y_position - (idx * 20), f"{bullet} {responsibility}")

    # Skills Header
    y_position = y_position - (len(responsibilities) * 20) - 40
    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, y_position, "Skills:")

    # Skills List
    c.setFont("Helvetica-Oblique", 12) 
    skills = llm_response.get("Skills")
    for idx, skill in enumerate(skills, 1):
        c.drawString(120, y_position - (idx * 20), f"{bullet} {skill}")

    # Experience Header
    y_position = y_position - (len(skills) * 20) - 40
    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, y_position, "Experience:")

    # Experience List
    c.setFont("Helvetica-Oblique", 12) 
    experiences = llm_response.get("Experience")
    for idx, ex in enumerate(experiences, 1):
        c.drawString(120, y_position - (idx * 20), f"{bullet} {ex}")

   
    y_position = y_position - (len(experiences) * 20) - 40 

    # Add an empty paragraph before the closing statement
    y_position-=20
    # Closing Statement at the Bottom Centered
    closing_y_position = y_position  # position the closing statement near the bottom
    c.setFont("Helvetica-Oblique", 12)
    text_width = c.stringWidth(llm_response.get("Closing Statement"), "Helvetica-Oblique", 12)
    c.drawString((width - text_width) / 2.0, closing_y_position, llm_response.get("Closing Statement"))

    # Save the PDF
    c.showPage()
    c.save()
    return f"{file_name}.pdf"


def save_jd_txt(llm_response, file_name):
    """Persist the JD as TXT"""
    save_jd_doc(llm_response, file_name)
    pypandoc.convert_file(
        f"{DOCS_FOLDER}/{file_name}.docx",
        "plain",
        outputfile=f"{DOCS_FOLDER}/{file_name}.txt",
    )
    return f"{file_name}.txt"

 # Helper function to wrap text within the specified width
def wrap_text(c,text, max_width):
    lines = []
    words = text.split(' ')
    current_line = ''
    for word in words:
        if c.stringWidth(current_line + word, "Helvetica", 12) < max_width:
            current_line += word + ' '
        else:
            lines.append(current_line.strip())
            current_line = word + ' '
    if current_line:
        lines.append(current_line.strip())
    return lines